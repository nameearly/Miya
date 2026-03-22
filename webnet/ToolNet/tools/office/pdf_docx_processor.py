"""
PDF和Word文档处理工具 - 弥娅核心模块
支持内容提取、文档分析
"""

import PyPDF2
from docx import Document
from typing import Dict, List, Optional, Any
from pathlib import Path
import re
import logging

logger = logging.getLogger(__name__)


class PDFDocxProcessor:
    """PDF和Word文档处理器"""

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']

    def load_file(self, file_path: str) -> Optional[Dict[str, Any]]:
        """加载PDF或Word文件"""
        path = Path(file_path)

        if not path.exists():
            logger.error(f"文件不存在: {file_path}")
            return None

        if path.suffix.lower() not in self.supported_formats:
            logger.error(f"不支持的文件格式: {path.suffix}")
            return None

        try:
            if path.suffix.lower() == '.pdf':
                return self._load_pdf(file_path)
            elif path.suffix.lower() == '.docx':
                return self._load_docx(file_path)
            else:
                logger.error(f"旧版.doc文件暂不支持，请转换为.docx")
                return None
        except Exception as e:
            logger.error(f"加载文件失败: {e}")
            return None

    def _load_pdf(self, file_path: str) -> Dict[str, Any]:
        """加载PDF文件"""
        reader = PyPDF2.PdfReader(file_path)
        metadata = reader.metadata

        # 提取所有页面文本
        full_text = ""
        page_contents = []

        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            page_contents.append({
                "页码": i + 1,
                "内容": page_text,
                "字符数": len(page_text)
            })
            full_text += page_text + "\n"

        return {
            "文件类型": "PDF",
            "总页数": len(reader.pages),
            "元数据": {
                "标题": metadata.get('/Title', ''),
                "作者": metadata.get('/Author', ''),
                "创建日期": metadata.get('/CreationDate', ''),
                "修改日期": metadata.get('/ModDate', ''),
            },
            "完整文本": full_text,
            "分页内容": page_contents,
            "总字符数": len(full_text),
            "总单词数": len(full_text.split())
        }

    def _load_docx(self, file_path: str) -> Dict[str, Any]:
        """加载Word文件"""
        doc = Document(file_path)

        # 提取段落文本
        paragraphs = []
        full_text = ""

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append({
                    "段落": i + 1,
                    "内容": text,
                    "字符数": len(text)
                })
                full_text += text + "\n"

        # 提取表格
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append({
                "表格": i + 1,
                "行数": len(table.rows),
                "列数": len(table.columns),
                "数据": table_data
            })

        return {
            "文件类型": "Word",
            "总段落数": len(paragraphs),
            "总表格数": len(tables),
            "完整文本": full_text,
            "段落内容": paragraphs,
            "表格内容": tables,
            "总字符数": len(full_text),
            "总单词数": len(full_text.split())
        }

    def extract_keywords(self, content: str, top_n: int = 10) -> List[Dict[str, Any]]:
        """从文本中提取关键词"""
        # 简单的中文分词和关键词提取
        # 这里使用基础的词频统计
        words = re.findall(r'[\u4e00-\u9fff]{2,}', content)

        word_freq = {}
        for word in words:
            # 过滤常见停用词
            if word not in ['的', '了', '是', '在', '和', '有', '我', '你', '他', '她']:
                word_freq[word] = word_freq.get(word, 0) + 1

        # 排序取前N个
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)

        keywords = []
        for word, freq in sorted_words[:top_n]:
            keywords.append({
                "关键词": word,
                "频次": freq,
                "重要度": freq / len(words) * 100
            })

        return keywords

    def analyze_document(self, file_path: str) -> Dict[str, Any]:
        """深度分析文档"""
        doc_data = self.load_file(file_path)

        if doc_data is None:
            return {"success": False, "error": "无法加载文档"}

        full_text = doc_data.get("完整文本", "")

        analysis = {
            "文件信息": doc_data,
            "关键词提取": self.extract_keywords(full_text, top_n=15),
            "文本统计": {
                "总字符数": len(full_text),
                "总行数": len(full_text.split('\n')),
                "平均句长": len(full_text) / len(full_text.split('。')) if full_text.split('。') else 0,
                "中文比例": len(re.findall(r'[\u4e00-\u9fff]', full_text)) / len(full_text) * 100
            },
            "结构分析": self._analyze_structure(doc_data)
        }

        return {"success": True, "analysis": analysis}

    def _analyze_structure(self, doc_data: Dict[str, Any]) -> Dict[str, Any]:
        """分析文档结构"""
        doc_type = doc_data.get("文件类型")

        if doc_type == "PDF":
            return {
                "类型": "PDF文档",
                "结构特点": [
                    f"共 {doc_data.get('总页数')} 页"
                ]
            }
        elif doc_type == "Word":
            return {
                "类型": "Word文档",
                "结构特点": [
                    f"包含 {doc_data.get('总段落数')} 个段落",
                    f"包含 {doc_data.get('总表格数')} 个表格"
                ],
                "表格信息": doc_data.get("表格内容", [])
            }

        return {}

    def extract_tables_from_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """从PDF中提取表格（简单实现）"""
        # 注意：PDF表格提取通常需要pdfplumber或tabula-py
        # 这里使用PyPDF2的简单实现
        import PyPDF2

        reader = PyPDF2.PdfReader(file_path)
        tables = []

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            # 简单的表格识别：连续的制表符或多个空格
            lines = text.split('\n')
            table_lines = []

            for line in lines:
                if '\t' in line or '    ' in line:
                    table_lines.append(line.split('\t') if '\t' in line else re.split(r'\s{3,}', line))

            if table_lines:
                tables.append({
                    "页码": page_num + 1,
                    "数据": table_lines
                })

        return tables

    def extract_images_from_docx(self, file_path: str, output_dir: str = "./extracted_images") -> List[str]:
        """从Word文档中提取图片"""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        doc = Document(file_path)
        image_paths = []

        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image = rel.target_part
                output_path = Path(output_dir) / Path(rel.target_ref).name
                output_path.parent.mkdir(parents=True, exist_ok=True)

                with open(output_path, "wb") as f:
                    f.write(image.blob)

                image_paths.append(str(output_path))
                logger.info(f"提取图片: {output_path}")

        return image_paths

    def search_in_document(self, file_path: str, keyword: str) -> List[Dict[str, Any]]:
        """在文档中搜索关键词"""
        doc_data = self.load_file(file_path)

        if doc_data is None:
            return []

        full_text = doc_data.get("完整文本", "")
        results = []

        # 简单的字符串搜索
        start = 0
        while True:
            pos = full_text.find(keyword, start)
            if pos == -1:
                break

            # 获取上下文（前后50个字符）
            context_start = max(0, pos - 50)
            context_end = min(len(full_text), pos + len(keyword) + 50)
            context = full_text[context_start:context_end]

            results.append({
                "位置": pos,
                "上下文": context,
                "关键词位置": pos - context_start
            })

            start = pos + 1

        logger.info(f"在文档中找到 {len(results)} 处 '{keyword}'")
        return results


def process_document_command(
    file_path: str,
    operations: List[str],
    output_path: str = None
) -> Dict[str, Any]:
    """
    执行文档处理命令

    Args:
        file_path: 输入文件路径
        operations: 操作列表
            ["analyze", "extract_tables", "search:关键词", "extract_images"]
        output_path: 输出路径（可选）

    Returns:
        处理结果
    """
    processor = PDFDocxProcessor()
    results = {
        "file_path": file_path,
        "operations": [],
        "results": {}
    }

    for op in operations:
        if op == "analyze":
            results["results"]["分析"] = processor.analyze_document(file_path)
            results["operations"].append("文档分析")

        elif op.startswith("search:"):
            keyword = op.split(":", 1)[1]
            search_results = processor.search_in_document(file_path, keyword)
            results["results"][f"搜索:{keyword}"] = search_results
            results["operations"].append(f"搜索关键词: {keyword}")

        elif op == "extract_tables":
            if file_path.endswith('.pdf'):
                tables = processor.extract_tables_from_pdf(file_path)
                results["results"]["表格"] = tables
                results["operations"].append("提取表格")
            else:
                results["results"]["表格"] = "PDF文件支持表格提取"
                results["operations"].append("尝试提取表格")

        elif op == "extract_images":
            if file_path.endswith('.docx'):
                images = processor.extract_images_from_docx(file_path, output_path or "./extracted_images")
                results["results"]["图片"] = images
                results["operations"].append("提取图片")
            else:
                results["results"]["图片"] = "仅Word文件支持图片提取"
                results["operations"].append("尝试提取图片")

    return results
